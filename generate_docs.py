"""
Génère les fichiers Word (.docx) depuis les fichiers Markdown.
Usage : python generate_docs.py
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


def style_doc(doc):
    """Applique les styles globaux au document."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for i in range(1, 5):
        h = doc.styles[f'Heading {i}']
        h.font.name = 'Calibri'
        h.font.bold = True
        if i == 1:
            h.font.size = Pt(20)
            h.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)
        elif i == 2:
            h.font.size = Pt(15)
            h.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
        elif i == 3:
            h.font.size = Pt(12)
            h.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
        elif i == 4:
            h.font.size = Pt(11)
            h.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)

    # Marges
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)


def parse_inline(text):
    """Retourne une liste de (texte, bold, italic, code)."""
    parts = []
    pattern = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], False, False, False))
        s = m.group()
        if s.startswith('`'):
            parts.append((s[1:-1], False, False, True))
        elif s.startswith('**') or s.startswith('__'):
            parts.append((s[2:-2], True, False, False))
        elif s.startswith('*'):
            parts.append((s[1:-1], False, True, False))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False, False, False))
    return parts


def add_inline(para, text):
    """Ajoute du texte inline avec formatage dans un paragraphe."""
    for content, bold, italic, code in parse_inline(text):
        run = para.add_run(content)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)


def md_to_docx(md_path, docx_path, title):
    doc = Document()
    style_doc(doc)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Page de titre
    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = RGBColor(0x19, 0x76, 0xD2)
    doc.add_paragraph()

    in_code = False
    code_lines = []
    in_table = False
    table_rows = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # ── Bloc de code ───────────────────────────────────────────
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
                    # Fond gris clair via XML
                    pPr = p._p.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:fill'), 'F3F3F3')
                    pPr.append(shd)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Tableau ────────────────────────────────────────────────
        if line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            # Ignorer les lignes de séparation |---|---|
            if all(re.match(r'^[-:]+$', c) for c in cells if c):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            # Vérifier si la ligne suivante est encore un tableau
            if i >= len(lines) or not lines[i].strip().startswith('|'):
                if table_rows:
                    ncols = max(len(r) for r in table_rows)
                    t = doc.add_table(rows=len(table_rows), cols=ncols)
                    t.style = 'Table Grid'
                    for ri, row in enumerate(table_rows):
                        for ci, cell_text in enumerate(row):
                            if ci < ncols:
                                c = t.cell(ri, ci)
                                c.text = ''
                                p = c.paragraphs[0]
                                add_inline(p, cell_text)
                                p.paragraph_format.space_before = Pt(3)
                                p.paragraph_format.space_after = Pt(3)
                                if ri == 0:
                                    p.runs[0].bold = True if p.runs else None
                                    set_cell_bg(c, '1976D2')
                                    for run in p.runs:
                                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                elif ri % 2 == 0:
                                    set_cell_bg(c, 'E3F2FD')
                    doc.add_paragraph()
                    table_rows = []
            continue

        # ── Ligne horizontale ──────────────────────────────────────
        if re.match(r'^-{3,}$', line.strip()) or re.match(r'^_{3,}$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Titres ─────────────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # Supprimer les emojis des titres pour Word
            h = doc.add_heading(text, level)
            h.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
            h.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ── Liste à puces ──────────────────────────────────────────
        m = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m:
            indent = len(m.group(1)) // 2
            text = m.group(2)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, text)
            i += 1
            continue

        # ── Liste numérotée ────────────────────────────────────────
        m = re.match(r'^\d+\.\s+(.*)', line)
        if m:
            text = m.group(1)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, text)
            i += 1
            continue

        # ── Bloc citation (>) ──────────────────────────────────────
        m = re.match(r'^>\s*(.*)', line)
        if m:
            text = m.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '10')
            left.set(qn('w:color'), '1976D2')
            pBdr.append(left)
            pPr.append(pBdr)
            add_inline(p, text)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)
            i += 1
            continue

        # ── Ligne vide ─────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Paragraphe normal ──────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_inline(p, line)
        i += 1

    doc.save(docx_path)
    print(f'OK : {docx_path}')


if __name__ == '__main__':
    import os
    base = os.path.dirname(os.path.abspath(__file__))

    md_to_docx(
        os.path.join(base, 'docs', 'guide_utilisateur.md'),
        os.path.join(base, 'docs', 'Guide_Utilisateur_MECA_AUTO.docx'),
        'MECA AUTO — Guide Utilisateur'
    )
    md_to_docx(
        os.path.join(base, 'DEPLOYMENT_RPI.md'),
        os.path.join(base, 'docs', 'Guide_Deploiement_RPi.docx'),
        'MECA AUTO — Guide de Déploiement Raspberry Pi'
    )
    md_to_docx(
        os.path.join(base, 'README.md'),
        os.path.join(base, 'docs', 'README_MECA_AUTO.docx'),
        'MECA AUTO — Documentation Projet'
    )
    print('\nFichiers Word generes dans le dossier docs/')
